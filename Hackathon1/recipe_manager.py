import requests
import psycopg2
import smtplib

CONNECTION = psycopg2.connect(host='localhost', user='postgres', password='1234', dbname='hackathon1')
CURSOR = CONNECTION.cursor()
EDAMAM_ID = '7c239491'
EDAMAM_KEY = 'b2450a48ca06718b544940eb54b7acc2'


class Recipe:
    def add_recipe_to_favorites(self, active_user, input_type):
        CURSOR.execute(f"SELECT recipe_name FROM favorites "  # finds user's favorites' names
                       f"INNER JOIN user_favorites ON user_favorites.recipe_id = favorites.recipe_id "
                       f"INNER JOIN users ON user_favorites.user_id = users.user_id AND users.username "
                       f"ILIKE '{active_user}';")
        # existing_faves = [i[0] for i in CURSOR.fetchall()] # not needed. users can have duplicate favorite names.
        # if self.name in existing_faves:
        #     if input_type == 'api':
        #         print(f"A favorite named {self.name} already exists. Pick another action.")
        #     elif input_type == 'user':
        #         print(f"A favorite named {self.name} already exists. Use a different name.")
        #         create_user_recipe(active_user)
        # else:
        CURSOR.execute(f"INSERT INTO favorites(recipe_name, ingredients, cook_time, cuisine, diet_code) VALUES"
                       f"('{self.name}', '{self.ingredients_string}', '{self.cook_time}', '{self.cuisine}', "
                       f"'{self.diet_code}');")
        CONNECTION.commit()
        CURSOR.execute(f"INSERT INTO user_favorites(user_id, recipe_id) VALUES"
                       f"((SELECT user_id FROM users WHERE username ILIKE '{active_user}'),"
                       f"(SELECT recipe_id FROM favorites WHERE recipe_id = "
                       f"(SELECT MAX(recipe_id) FROM favorites)))")
        CONNECTION.commit()
        print(f"\n{self.name} added to Favorites")


class APIRecipe(Recipe):
    def __init__(self, json_data):
        self.name = json_data['recipe']['label'].replace("'", '')
        self.ingredients = json_data['recipe']['ingredientLines']
        self.ingredients_string = ', '.join(self.ingredients).replace("'", '')
        self.cook_time = int(json_data['recipe']['totalTime'])
        self.cuisine = json_data['recipe']['cuisineType'][0].replace("'", '')
        self.diet_code = 3 if 'Celiac' in json_data['recipe']['healthLabels'] else \
            2 if 'Vegan' in json_data['recipe']['healthLabels'] else \
                1 if 'Vegetarian' in json_data['recipe']['healthLabels'] else 0


class UserRecipe(Recipe):
    def __init__(self):
        # input of new recipe details
        self.name = input("\nRecipe Name: ").replace("'", '').title()
        while True:
            self.cook_time = input("Cook time [min]: ")
            try:
                self.cook_time = int(self.cook_time)
            except ValueError:
                print("Enter an integer!")
                continue
            else:
                break
        self.cuisine = input("Cuisine: ").replace("'", "")
        while True:
            self.diet_code = input("Enter diet code. 0: None, 1: Vegetarian, 2: Vegan, 3: Celiac ")
            if self.diet_code not in ['0', '1', '2', '3']:
                print("Choose 0, 1, 2, or 3!")
                continue
            else:
                break
        # ingredient entry:
        self.ingredients = []
        done = False
        while not done:
            ingredient = input("Enter ingredient. Type 'done' when you're done: ")
            if ingredient.lower() == 'done':
                break
            else:
                self.ingredients.append(ingredient)
        self.ingredients_string = ', '.join(self.ingredients).replace("'", "")


def get_search_parameters(active_user):
    params = {}
    valid_kw = False
    while not valid_kw:
        search_keywords = ',+'.join(input("\nKeywords to search (with spaces): ").split())
        if search_keywords != '':
            break
    cuisine_input = input('Enter cuisine: ')
    cuisine = f"&cuisineType={cuisine_input}" if cuisine_input != '' else ''
    diet_input = input('Enter diet restriction: ')
    diet = f"&cuisineType={diet_input}" if diet_input != '' else ''
    params['search_keywords'] = search_keywords
    params['diet'] = diet
    params['cuisine'] = cuisine
    search_all_recipes(params, active_user)


def search_all_recipes(params, active_user):
    URL = f"https://api.edamam.com/search?q={params['search_keywords']}&app_id={EDAMAM_ID}&app_key={EDAMAM_KEY}" \
          f"{params['diet']}{params['cuisine']}"
    search_results = requests.get(url=URL).json()['hits']
    titles = [i['recipe']['label'] for i in search_results]
    print(titles)
    if len(titles) == 0:
        print("No results. Search again.")
        get_search_parameters(active_user)
    enum = enumerate(titles)
    for count, value in enum:
        print(f"{count + 1}.", value)
    while True:
        view_choice = input("Which # recipe would you like to view? (1-10): ")
        try:
            view_choice = int(view_choice)
        except ValueError:
            print("Pick a number!")
            continue
        else:
            if 1 <= view_choice <= 10:
                a = APIRecipe(search_results[view_choice - 1])
                view_specific_api_recipe(a, params, active_user)
                return
            else:
                print("Number must be 1-10!")
                continue


def view_favorites(active_user):
    CURSOR.execute(f"SELECT recipe_name FROM favorites "  # finds user's favorites' names
                   f"INNER JOIN user_favorites ON user_favorites.recipe_id = favorites.recipe_id "
                   f"INNER JOIN users ON user_favorites.user_id = users.user_id AND users.username "
                   f"ILIKE '{active_user}';")
    recipe_response = [i[0] for i in CURSOR.fetchall()]
    enum = enumerate(recipe_response)
    print("\nFavorites:")
    for count, value in enum:
        print(f"{count + 1}.", value)
    view_choice = input("Which recipe would you like to view? Or (B) to go back: ").lower()
    if view_choice not in [str(int(i + 1)) for i in list(range(len(recipe_response)))] + ['b']:
        print("\nChoose a relevant number or (B)!")
        view_favorites(active_user)
    elif view_choice == 'b':
        return
    else:
        chosen_favorite = recipe_response[int(view_choice) - 1]
        print(chosen_favorite)
        view_specific_favorite(active_user, view_choice)


def view_specific_api_recipe(recipe, params, active_user):
    CURSOR.execute(f"SELECT diet_name FROM diets WHERE diet_code = {recipe.diet_code}")  # converts API's diet info
    diet_type = CURSOR.fetchall()[0][0]
    print(f"\n{recipe.name}")
    print(f"Cuisine: {recipe.cuisine.title()}")
    print(f"Diet Type: {diet_type}")
    ingredients_print = recipe.ingredients_string.replace('\n', ", ")
    print(f"\nIngredients:\n{ingredients_print}")
    while True:
        recipe_view_choice = input("\nWhat would you like to do?\n(E)mail recipe\n(A)dd to favorites\n"
                                   "Back to search (R)esults\n(H)ome: ")
        if recipe_view_choice.lower() == 'a':
            recipe.add_recipe_to_favorites(active_user, 'api')
        elif recipe_view_choice.lower() == 'e':
            recipe.email_recipe(active_user)
        elif recipe_view_choice.lower() == 'r':
            search_all_recipes(params, active_user)
        elif recipe_view_choice.lower() == 'h':
            break


def view_specific_favorite(active_user, view_choice):
    CURSOR.execute(f"SELECT favorites.recipe_name, favorites.cuisine, favorites.diet_code, "
                   f"favorites.cook_time,"
                   f" favorites.ingredients, user_favorites.recipe_id FROM favorites "  # finds user's favorites
                   f"INNER JOIN user_favorites ON user_favorites.recipe_id = favorites.recipe_id "
                   f"INNER JOIN users ON user_favorites.user_id = users.user_id AND users.username "
                   f"ILIKE '{active_user}';")
    fav_info = CURSOR.fetchall()[int(view_choice) - 1]
    recipe_name = fav_info[0]
    cuisine = fav_info[1]
    diet = fav_info[2]
    cook_time = fav_info[3]
    ingredients = fav_info[4]
    recipe_id = fav_info[5]
    ingredients_print = ingredients.replace(', ', "\n")
    display_text = f"\n{recipe_name}\nCuisine: {cuisine.title()}\nDiet type: {diet}\nCook time: {cook_time} minutes\n" \
                   f"\nIngredients:\n{ingredients_print}"
    print(display_text)
    back_email_or_delete = input("(B)ack, (E)mail favorite, or (D)elete this favorite? ")
    if back_email_or_delete.upper() == 'D':
        CURSOR.execute(f"DELETE FROM favorites WHERE recipe_id = {recipe_id}")
        CONNECTION.commit()
    elif back_email_or_delete.upper() == 'E':
        email_recipe(recipe_name, display_text, active_user)


def create_user_recipe(active_user):
    ur = UserRecipe()
    ur.add_recipe_to_favorites(active_user, 'user')


def email_recipe(recipe_name, display_text, active_user):
    from email.mime.application import MIMEApplication
    from email.mime.base import MIMEBase
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email import encoders
    import docx
    import os
    # from docx import Document
    # from docx.shared import Inches
    document = docx.Document()
    document.add_heading(recipe_name, 0)
    document.add_paragraph(display_text)
    document.save(f"{recipe_name}.docx")

    from passwords import gmail_pw
    SENDER = 'akivabuckman@gmail.com'
    RECEIVER = 'akivabuckman@yahoo.com'

    msg = MIMEMultipart()
    msg['From'] = SENDER
    msg['To'] = RECEIVER
    msg['Subject'] = recipe_name
    body = f"{recipe_name} recipe attached."
    msg.attach(MIMEText(body))
    attachment = open(fr"C:\Users\akiva\Documents\Coding\DI\DI_Bootcamp\Hackathon1\{recipe_name}.docx", 'rb')
    part = MIMEBase("application", "octet-stream")
    part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition",
                    f"attachment; filename= {recipe_name}.docx")
    msg.attach(part)

    msg = msg.as_string()
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(SENDER, gmail_pw)
        smtp.sendmail(SENDER, RECEIVER, msg)
