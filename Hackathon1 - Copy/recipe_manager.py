import requests
import psycopg2
import smtplib
import tkinter as tk
import recipe_objects
from passwords.passwords import gmail_pw

CONNECTION = psycopg2.connect(host='localhost', user='postgres', password='1234', dbname='hackathon1')
CURSOR = CONNECTION.cursor()
EDAMAM_ID = '7c239491'
EDAMAM_KEY = 'b2450a48ca06718b544940eb54b7acc2'


class RecipeManager:
    def run_choice(self, function, active_user):
        function(self, active_user)

    def get_search_parameters(self, active_user):
        self.search_screen = tk.Tk()
        self.search_screen.title("Search for a Recipe")
        label1 = tk.Label(text="Enter search keywords:")
        label1.grid(row=0, column=0, padx=5, pady=5)
        label1.focus()
        self.keyword_entry = tk.Entry(self.search_screen)
        self.keyword_entry.grid(row=0, column=1, padx=5, pady=5)
        label2 = tk.Label(text="Cuisine:")
        label2.grid(row=1, column=0, padx=5, pady=5)
        self.cuisine_entry = tk.Entry(self.search_screen)
        self.cuisine_entry.grid(row=1, column=1, padx=5, pady=5)
        label3 = tk.Label(text="Diet restriction:")
        label3.grid(row=2, column=0)
        self.diet_entry = tk.Entry(self.search_screen)
        self.diet_entry.grid(row=2, column=1, padx=5, pady=5)
        search_button = tk.Button(text="Search", command=lambda: self.validate_and_search(active_user))
        search_button.grid(row=3, column=0)
        back_button = tk.Button(text="Home", command=lambda: self.return_now(self.search_screen))
        back_button.grid(row=3, column=1)
        self.search_screen.mainloop()

    def validate_and_search(self, active_user):
        self.search_keywords = ',+'.join(self.keyword_entry.get().split())
        self.search_screen_error = tk.Label(text="", fg='red')
        self.search_screen_error.grid(row=4, column=0, columnspan=2)

        if self.search_keywords == '':
            self.search_screen_error.config(text="Keywords can't be empty!")
        else:
            cuisine = f"&cuisineType={self.cuisine_entry.get()}" if self.cuisine_entry.get() != '' else ''
            diet = f"&cuisineType={self.diet_entry.get()}" if self.diet_entry.get() != '' else ''
            params = {}
            params['search_keywords'] = self.search_keywords
            params['diet'] = diet
            params['cuisine'] = cuisine
            self.search_recipe_api(params, active_user, self.search_screen)

    def search_recipe_api(self, params, active_user, previous_screen):
        URL = f"https://api.edamam.com/search?q={params['search_keywords']}&app_id={EDAMAM_ID}&app_key={EDAMAM_KEY}" \
              f"{params['diet']}{params['cuisine']}"
        self.search_results = requests.get(url=URL).json()['hits']
        self.titles = [i['recipe']['label'] for i in self.search_results]
        if len(self.titles) == 0:
            self.search_screen_error.config(text="No results. Change search.")
        else:
            self.display_api_results(active_user, params, previous_screen)

    def display_api_results(self, active_user, params, previous_screen):
        previous_screen.destroy()
        self.api_results_screen = tk.Tk()
        self.api_results_screen.title("Search Results")
        label1 = tk.Label(text=f"Results for {self.search_keywords.replace('+', ' ')}:")
        label1.grid(row=0, column=0, columnspan=2, pady=5)
        for count, value in enumerate(self.titles):
            label = tk.Label(text=f"{count + 1}. {value}")
            label.grid(row=count + 1, column=0, padx=5, pady=5)
            json_data = self.search_results[count]
            self.create_api_recipe_button(count, active_user, json_data, params)
        home_button = tk.Button(text="Back", command=lambda: self.return_now(self.api_results_screen))
        home_button.grid(row=len(self.titles) + 2, column=0, columnspan=2)

    def return_now(self, current_screen):
        current_screen.destroy()

    def create_api_recipe_button(self, count, active_user, json_data, params):
        new_button = tk.Button(text="View Recipe",
                               command=lambda: self.view_specific_api_recipe2(json_data, active_user, params))
        new_button.grid(row=1 + count, column=1, padx=5, pady=5)

    def view_specific_api_recipe2(self, json_data, active_user, params):
        apir = recipe_objects.APIRecipe(json_data)
        CURSOR.execute(f"SELECT diet_name FROM diets WHERE diet_code = {apir.diet_code}")
        diet_type = CURSOR.fetchall()[0][0]
        self.api_results_screen.destroy()
        self.specific_api_screen = tk.Tk()
        self.specific_api_screen.title(apir.name)
        label1 = tk.Label(text=apir.name)
        label1.grid(row=0, column=0)
        label2 = tk.Label(text=f"Cuisine: {apir.cuisine}")
        label2.grid(row=1, column=0)
        label3 = tk.Label(text=f"Diet Type: {diet_type}")
        label3.grid(row=2, column=0)
        ingredients_display = apir.ingredients_string.replace(", ", '\n')
        label4 = tk.Label(text=f"Ingredients:\n{ingredients_display}")
        label4.grid(row=3, column=0)
        add_to_favorites_button = tk.Button(text="Add to Favorites",
                                            command=lambda: apir.add_recipe_to_favorites(active_user,
                                                                                         self.specific_api_screen,
                                                                                         add_to_favorites_button))
        add_to_favorites_button.grid(row=4, column=0, pady=5)
        back_to_results_button = tk.Button(text="Back to Results",
                                           command=lambda: self.search_recipe_api(params, active_user,
                                                                                  self.specific_api_screen))
        back_to_results_button.grid(row=5, column=0, pady=5)
        back_home_button = tk.Button(text="Home", command=lambda: self.return_now(self.specific_api_screen))
        back_home_button.grid(row=6, column=0, pady=5)

    # def view_specific_api_recipe(recipe, params, active_user):
    #     CURSOR.execute(f"SELECT diet_name FROM diets WHERE diet_code = {recipe.diet_code}")  # converts API's diet info
    #     diet_type = CURSOR.fetchall()[0][0]
    #     print(f"\n{recipe.name}")
    #     print(f"Cuisine: {recipe.cuisine.title()}")
    #     print(f"Diet Type: {diet_type}")
    #     ingredients_print = recipe.ingredients_string.replace('\n', ", ")
    #     print(f"\nIngredients:\n{ingredients_print}")
    #     while True:
    #         recipe_view_choice = input("\nWhat would you like to do?\n(A)dd to favorites\n"
    #                                    "Back to search (R)esults\n(H)ome: ")
    #         if recipe_view_choice.lower() == 'a':
    #             recipe.add_recipe_to_favorites(active_user, 'api')
    #         elif recipe_view_choice.lower() == 'r':
    #             search_all_recipes(params, active_user)
    #         elif recipe_view_choice.lower() == 'h':
    #             break

    # def search_all_recipes(params, active_user):
    #     URL = f"https://api.edamam.com/search?q={params['search_keywords']}&app_id={EDAMAM_ID}&app_key={EDAMAM_KEY}" \
    #           f"{params['diet']}{params['cuisine']}"
    #     search_results = requests.get(url=URL).json()['hits']
    #     titles = [i['recipe']['label'] for i in search_results]
    #     print(titles)
    #     if len(titles) == 0:
    #         print("No results. Search again.")
    #         get_search_parameters(active_user)
    #     enum = enumerate(titles)
    #     for count, value in enum:
    #         print(f"{count + 1}.", value)
    #     while True:
    #         view_choice = input("Which # recipe would you like to view? (1-10): ")
    #         try:
    #             view_choice = int(view_choice)
    #         except ValueError:
    #             print("Pick a number!")
    #             continue
    #         else:
    #             if 1 <= view_choice <= 10:
    #                 a = APIRecipe(search_results[view_choice - 1])
    #                 view_specific_api_recipe(a, params, active_user)
    #                 return
    #             else:
    #                 print("Number must be 1-10!")
    #                 continue

    def view_favorites(self, active_user):
        CURSOR.execute(f"SELECT recipe_name FROM favorites "  # finds user's favorites' names
                       f"INNER JOIN user_favorites ON user_favorites.recipe_id = favorites.recipe_id "
                       f"INNER JOIN users ON user_favorites.user_id = users.user_id AND users.username "
                       f"ILIKE '{active_user}';")
        recipe_response = [i[0] for i in CURSOR.fetchall()]
        self.favorite_list_screen = tk.Tk()
        self.favorite_list_screen.title(f"{active_user.title()}'s Favorites")
        label = tk.Label(text=f"{active_user.title()}'s Favorites:")
        label.grid(row=0, column=0, columnspan=2)
        for count, value in enumerate(recipe_response):
            label = tk.Label(text=f"{count + 1}. {value}")
            label.grid(row=count + 1, column=0)
            self.create_favorite_button(count=count, value=value, active_user=active_user, params=1)

        self.favorite_list_screen.mainloop()
        #     print(f"{count + 1}.", value)
        # view_choice = input("Which recipe would you like to view? Or (B) to go back: ").lower()
        # if view_choice not in [str(int(i + 1)) for i in list(range(len(recipe_response)))] + ['b']:
        #     print("\nChoose a relevant number or (B)!")
        #     view_favorites(active_user)
        # elif view_choice == 'b':
        #     return
        # else:
        #     chosen_favorite = recipe_response[int(view_choice) - 1]
        #     print(chosen_favorite)
        #     view_specific_favorite(active_user, view_choice)

    def create_favorite_button(self, count, value, active_user, params):
        new_button = tk.Button(text="View Recipe",
                               # command=lambda: self.view_specific_fav_recipe(json_data, active_user, params))
                               command=lambda: self.view_specific_fav_recipe(active_user, count))

        new_button.grid(row=1 + count, column=1, padx=5, pady=5)

    def view_specific_fav_recipe(self, active_user, count):
        CURSOR.execute(f"SELECT favorites.recipe_name, favorites.cuisine, favorites.diet_code, "
                       f"favorites.cook_time,"
                       f" favorites.ingredients, user_favorites.recipe_id FROM favorites "  # finds user's favorites
                       f"INNER JOIN user_favorites ON user_favorites.recipe_id = favorites.recipe_id "
                       f"INNER JOIN users ON user_favorites.user_id = users.user_id AND users.username "
                       f"ILIKE '{active_user}';")
        fav_info = CURSOR.fetchall()[count]
        recipe_name = fav_info[0]
        cuisine = fav_info[1]
        diet = fav_info[2]
        cook_time = fav_info[3]
        ingredients = fav_info[4]
        recipe_id = fav_info[5]
        ingredients_print = ingredients.replace(', ', "\n")
        display_text = f"\n{recipe_name}\nCuisine: {cuisine.title()}\nDiet type: {diet}\nCook time: {cook_time} minutes\n" \
                       f"\nIngredients:\n{ingredients_print}"
        self.return_now(self.favorite_list_screen)
        specific_fav_screen = tk.Tk()
        specific_fav_screen.title(recipe_name)
        label1 = tk.Label(text=recipe_name.title())
        label1.grid(row=0, column=0, columnspan=2, padx=5, pady=5)
        label2 = tk.Label(text=f"Cuisine: {cuisine}")
        label2.grid(row=1, column=0, columnspan=2, padx=5, pady=5)
        label3 = tk.Label(text=f"Diet: {diet}")
        label3.grid(row=2, column=0, columnspan=2, padx=5, pady=5)
        label4 = tk.Label(text=f"Cook Time: {cook_time} min")
        label4.grid(row=3, column=0, columnspan=2, padx=5, pady=5)
        label5 = tk.Label(text="Ingredients:")
        label5.grid(row=4, column=0, columnspan=2, pady=5, padx=5)
        label6 = tk.Label(text=ingredients_print)
        label6.grid(row=5, column=0, columnspan=2, pady=5, padx=5)
        email_button = tk.Button(text="Email Recipe",
                                 command=lambda: self.email_recipe(recipe_name, display_text, active_user))
        email_button.grid(row=6, column=0, padx=5, pady=5)
        delete_button = tk.Button(text="Unfavorite", command=lambda: self.unfavorite(recipe_id))
        delete_button.grid(row=6, column=1, padx=5, pady=5)

    def unfavorite(self, recipe_id):
        CURSOR.execute(f"DELETE FROM favorites WHERE recipe_id = {recipe_id}")
        CONNECTION.commit()
    def create_user_recipe(self, active_user):
        ur = recipe_objects.UserRecipe(active_user)

    def create_api_recipe(self, active_user, json_data, params):
        apir = recipe_objects.APIRecipe(json_data)
        self.view_specific_api_recipe2(apir, json_data, active_user, params)

    def email_recipe(self, recipe_name, display_text, active_user):
        from email.mime.application import MIMEApplication
        from email.mime.base import MIMEBase
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email import encoders
        import docx
        import os
        # from docx import Document
        # from docx.shared import Inches
        document = docx.Document()
        document.add_heading(recipe_name, 0)
        document.add_paragraph(display_text)
        document.save(f"saved_files/{recipe_name}.docx")

        SENDER = 'akivabuckman@gmail.com'
        RECEIVER = 'akivabuckman@yahoo.com'

        msg = MIMEMultipart()
        msg['From'] = SENDER
        msg['To'] = RECEIVER
        msg['Subject'] = recipe_name
        body = f"Hey {active_user},\n{recipe_name} recipe attached."
        msg.attach(MIMEText(body))
        attachment = open(
            fr"C:\Users\akiva\Documents\Coding\DI\DI_Bootcamp\Hackathon1 - Copy\saved_files\{recipe_name}.docx", 'rb')
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition",
                        f"attachment; filename= {recipe_name}.docx")
        msg.attach(part)

        msg = msg.as_string()
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(SENDER, gmail_pw)
            smtp.sendmail(SENDER, RECEIVER, msg)
